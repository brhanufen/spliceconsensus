"""Build the docx report (mirrors SpliceConsensus_Report.tex / README.md). Run from repo root."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "SpliceConsensus_Report.docx"
AC = RGBColor(0x1A, 0x52, 0x76)
d = Document(); d.styles["Normal"].font.name = "Calibri"; d.styles["Normal"].font.size = Pt(11)

def h(t, l=1):
    p = d.add_heading(t, l)
    for r in p.runs: r.font.color.rgb = AC
    return p
def para(t): d.add_paragraph(t)

d.add_heading("A Reproducible MFASS Benchmark of Splice-Disruption Predictors Reveals a Shared Exon-Interior Blind Spot", 0)
s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Brhanu F. Znabu (1) and Zohaib Atif (2). June 2026."); r.italic = True; r.font.size = Pt(10)
aff = d.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = aff.add_run("(1) Biomedical Engineering Program, College of Engineering, University of Nebraska-Lincoln, Lincoln, NE, USA.   "
                 "(2) Department of Biomedical Science and Engineering, Gwangju Institute of Science and Technology (GIST), Gwangju, Republic of Korea."); ar.italic = True; ar.font.size = Pt(9)

h("Abstract")
para("We benchmark three published splicing variant-effect predictors, run through the Proto tool ecosystem, "
     "against a multiplexed experimental splicing assay. On 27,733 single-nucleotide variants in and around human exons from MFASS "
     "with measured exon-inclusion outcomes, Pangolin is the strongest predictor of splice-disrupting variants "
     "(AUROC 0.888, average precision 0.421), ahead of SpliceAI (0.819, 0.321) and SpliceTransformer (0.786, 0.317), "
     "and all three clearly exceed the older SPANR model (0.748, 0.228). The ranking reproduces the relative "
     "performance reported by the Pangolin authors, a correctness check on the pipeline. A calibrated consensus of "
     "the three modern predictors, evaluated on an exon-grouped held-out split, does not meaningfully improve over "
     "Pangolin alone. Stratifying by distance to the splice site exposes a shared blind spot: all four tools detect "
     "disruptions within a few bases of the splice site well, but recall is substantially lower in the exon interior, and 22% of "
     "disrupting variants are missed by every tool; these shared misses are enriched among variants away from splice sites, and many are exon-interior. Every number is computed against fixed "
     "experimental ground truth and is reproducible from the public dataset and the released code.")

h("Introduction")
para("RNA splicing is disrupted in a large share of disease-causing variants, and splice-altering variants are among the "
     "hardest to interpret in clinical genetics, because their effect depends on sequence context that is not obvious from "
     "the variant alone (Scotti and Swanson 2016; Wang and Cooper 2007; Cartegni et al. 2002). Deep-learning predictors of splicing from primary sequence, notably "
     "SpliceAI (Jaganathan et al. 2019) and Pangolin (Zeng and Li 2022), are now widely used to flag candidate "
     "splice-disrupting variants and appear in clinical variant-interpretation pipelines (Richards et al. 2015; Walker et al. 2023).")
para("These predictors, together with the more recent SpliceTransformer (You et al. 2024), the earlier SPANR model "
     "(Xiong et al. 2015), and other models such as MMSplice (Cheng et al. 2019), are usually evaluated in their own publications on heterogeneous datasets and metrics, which "
     "makes it hard to compare them on equal footing or to know where each is reliable. Multiplexed experimental assays "
     "of splicing provide a fixed ground truth for such comparisons (Rosenberg et al. 2015): MFASS (Chong et al. 2019) measured the splicing effect of tens of "
     "thousands of human single-nucleotide variants in exons and flanking intronic regions in a minigene reporter and labels those that strongly reduce exon "
     "inclusion as splice-disrupting variants.")
para("Prior work has benchmarked deep-learning splice predictors against functional splicing assays (Riepe et al. 2021). "
     "Here we provide a uniform, reproducible, open benchmark of all four predictors on the MFASS single-nucleotide "
     "variants, scored through a single interface, with bootstrap confidence intervals, an exon-grouped held-out consensus, and a legacy "
     "baseline. Beyond the aggregate ranking, we ask where the predictions fail: we stratify detection by distance to the "
     "splice site and identify the variants that every tool misses. The contribution is the reproducible benchmark and a "
     "characterization of a shared exon-interior blind spot.")

h("Data and ground truth")
para("MFASS (the Multiplexed Functional Assay of Splicing by Sort-seq; Chong et al. 2019) measured the splicing "
     "effect of human single-nucleotide variants in exons and flanking intronic regions in a minigene reporter. Each variant carries a measured change "
     "in exon inclusion, and variants reducing inclusion by at least 0.50 are labeled splice-disrupting variants (SDVs). "
     "We score the 28,972 mutant SNVs (51% in exons, 49% in the flanking introns within about 50 bp of a splice site); for the binary detection task we drop variants whose inclusion change was not "
     "measured, leaving 27,733 variants of which 1,050 (3.79%) are SDVs. The strong imbalance makes average precision "
     "the primary metric (Saito and Rehmsmeier 2015), reported alongside AUROC.")

h("Predictors and scoring")
para("Three predictors are run through Proto's wrappers: SpliceAI (Jaganathan et al. 2019), a deep convolutional model; "
     "Pangolin (Zeng and Li 2022), a tissue-aware successor; and SpliceTransformer (You et al. 2024), a transformer "
     "predicting tissue-specific splicing. The older SPANR model (Xiong et al. 2015) is included as a legacy baseline. "
     "Each variant is scored in its real hg38 context. For each tool we reduce its output to a single disruption-magnitude "
     "delta score. Tissue-aware tools are used in their tissue-agnostic setting, since MFASS is an in-vitro minigene assay.")

h("Results")
tbl = d.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
for i, c in enumerate(["predictor", "AUROC", "AP (AUPRC)", "AP 95% CI"]): tbl.rows[0].cells[i].text = c
for row in [["Pangolin", "0.888", "0.421", "[0.386, 0.454]"], ["SpliceAI", "0.819", "0.321", "[0.294, 0.353]"],
            ["SpliceTransformer", "0.786", "0.317", "[0.285, 0.352]"], ["SPANR (legacy)", "0.748", "0.228", "[0.201, 0.258]"]]:
    cells = tbl.add_row().cells
    for i, v in enumerate(row): cells[i].text = v
para("")
for fig, cap in [("results/fig_roc.png", "ROC for splice-disrupting-variant detection. Pangolin shows the strongest ROC performance; all modern tools exceed the legacy SPANR baseline."),
                 ("results/fig_pr.png", "Precision-recall for the same task (dashed line is the 0.038 positive base rate).")]:
    d.add_picture(str(ROOT / fig), width=Inches(4.6)); cp = d.add_paragraph(); cr = cp.add_run(cap); cr.italic = True; cr.font.size = Pt(9)
para("Finding 1: Pangolin is the strongest predictor, by a clear margin on both metrics; its AP confidence interval does "
     "not overlap the others. Finding 2: all three modern predictors beat the legacy SPANR model. Finding 3: the ranking "
     "reproduces the ordering the Pangolin authors reported, a correctness check on the pipeline.")
para("Finding 4 (output granularity): in these outputs SpliceAI's scores are more discretized than Pangolin's. SpliceAI reports its delta to two decimals and "
     "assigns exactly zero to most variants away from splice sites: 55% of non-disrupting variants but only 14% of SDVs "
     "score zero, so the zeros are genuine no-effect calls, not an artifact. Pangolin instead uses a continuous range, "
     "taking 17,723 distinct score values to SpliceAI's 101, with no variant scored exactly zero, and this finer score resolution may contribute to its stronger ranking here.")
para("Finding 5 (consensus): combining does not help. A logistic-regression consensus of the three modern predictors, "
     "trained and evaluated on an exon-grouped 50/50 split, reaches held-out AP 0.443 versus Pangolin's 0.442 on the same "
     "split, a difference far inside the bootstrap interval, and a slightly lower AUROC. The fitted model essentially "
     "relies on Pangolin alone. For this MFASS benchmark, Pangolin is the strongest single predictor.")

h("Where the tools fail")
para("The aggregate ranking hides where the predictions break down. We stratified splice-disrupting-variant recall by the "
     "distance of each variant to the nearest splice site, computed from the MFASS minigene construct annotation and "
     "cross-checked against an independent genomic-coordinate distance (Pearson r=0.997), comparing all four tools at a "
     "common operating point: the score threshold that gives a 10% false-positive rate on non-disrupting variants.")
d.add_picture(str(ROOT / "results/fig_failuremode.png"), width=Inches(5.4))
cp=d.add_paragraph(); cr=cp.add_run("Recall of splice-disrupting variants, at a common 10% false-positive rate, as a function "
     "of distance to the nearest splice site. Every tool detects splice-site-proximal disruptions well and degrades with distance "
     "from the splice site."); cr.italic=True; cr.font.size=Pt(9)
para("Recall declines with distance from the splice site for every tool. Within 2 bp of a splice site the modern tools "
     "recover 80 to 90% of disrupting variants (Pangolin 0.90, SpliceAI 0.82, SpliceTransformer 0.80); beyond 20 bp their "
     "recall falls to 0.48, 0.35, and 0.29 respectively, and SPANR to 0.20. More than half of the disrupting variants (53%) "
     "lie more than 10 bp from a splice site, and even Pangolin, the best tool, detects only 59% of them.")
para("Most strikingly, 22% of all disrupting variants (228 of 1,050) are missed by every one of the four tools at this "
     "operating point. These shared misses sit predominantly away from the splice site: 75% lie more than 10 bp from a splice "
     "site, and 162 of 228 (71%) are exonic. They are consistent with exon-interior regulatory disruptions, including possible exonic splicing enhancer or silencer effects (Cartegni et al. 2002), that splice-site models "
     "are not built to capture, the class of variant MFASS was designed to expose. The practical implication is that these "
     "tools are strongest for splice-site-proximal variants but require caution for exon-interior variants, where roughly one "
     "in five disrupting variants is invisible to all current predictors.")

h("What this shows, and what it does not")
para("It shows that, against a multiplexed experimental splicing assay, Pangolin detects splice-disrupting variants "
     "better than SpliceAI, SpliceTransformer, and SPANR, that all modern tools beat the legacy baseline, and that a "
     "simple consensus does not improve on the best single tool. Caveats: the predictors were trained on genomic and "
     "transcriptomic data and may have encountered these exons' wild-type splice sites during training, so this is a "
     "realistic benchmark rather than a fully held-out one (the same for all tools); MFASS is a single-context in-vitro "
     "assay, so the tissue-aware tools were used tissue-agnostically; and all predictors are scored as disruption "
     "magnitudes against a loss-of-function label. The contribution is a reproducible head-to-head benchmark, not a claim of clinical validation.")

h("Reproducibility and data processing")
para("The labels and continuous readout are the public MFASS measurements; all predictor scores are produced by "
     "the public models run through Proto. Minus-strand alleles were normalized to the "
     "plus strand before scoring; for the three Proto-scored modern predictors, full coverage of 28,972 of 28,972 variants confirms no variant was silently dropped (SPANR covers slightly fewer, a precomputed subset). SpliceAI "
     "returns no score for variants outside an annotated gene; these were assigned a delta of zero, a minority as shown "
     "by the depletion of disrupting variants among zero-scored variants. Every number is reproduced by score_one_tool.py "
     "(scoring) and analyze.py (metrics, consensus, and figures) from the public dataset.")

h("Data and code availability")
para("All code, the per-tool variant scores, the benchmark and failure-mode outputs, and the figures are openly available "
     "at github.com/brhanufen/spliceconsensus and archived on Zenodo (doi:10.5281/zenodo.20948820). The benchmark and all "
     "figures reproduce on a CPU in about a minute from the committed scores and a slim labels file, with no large download. "
     "The MFASS dataset is from Chong et al. 2019 (github.com/KosuriLab/MFASS); the predictors are the public SpliceAI, "
     "Pangolin, SpliceTransformer, and SPANR models, run through the Proto tool ecosystem.")

h("Author contributions, competing interests, and funding")
para("B.F.Z. designed the benchmark, ran the analyses, generated the figures, and wrote the manuscript. Z.A. contributed "
     "to the conceptual discussion and revised the manuscript. Both authors read and approved the final manuscript. The authors "
     "declare no competing interests. No specific funding was received for this work.")

h("References")
for ref in [
    "Scotti, M.M. and Swanson, M.S. RNA mis-splicing in disease. Nature Reviews Genetics 17(1):19–32, 2016. doi:10.1038/nrg.2015.3.",
    "Wang, G.-S. and Cooper, T.A. Splicing in disease: disruption of the splicing code and the decoding machinery. Nature Reviews Genetics 8:749–761, 2007. doi:10.1038/nrg2164.",
    "Cartegni, L., Chew, S.L., and Krainer, A.R. Listening to silence and understanding nonsense: exonic mutations that affect splicing. Nature Reviews Genetics 3:285–298, 2002. doi:10.1038/nrg775.",
    "Jaganathan, K., et al. Predicting splicing from primary sequence with deep learning (SpliceAI). Cell 176(3):535–548.e24, 2019. doi:10.1016/j.cell.2018.12.015.",
    "Zeng, T. and Li, Y.I. Predicting RNA splicing from DNA sequence using Pangolin. Genome Biology 23:103, 2022. doi:10.1186/s13059-022-02664-4.",
    "Richards, S., Aziz, N., Bale, S., et al. Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology. Genetics in Medicine 17:405–424, 2015. doi:10.1038/gim.2015.30.",
    "Walker, L.C., de la Hoya, M., Wiggins, G.A.R., et al. Using the ACMG/AMP framework to capture evidence related to predicted and observed impact on splicing. The American Journal of Human Genetics 110:1046–1067, 2023. doi:10.1016/j.ajhg.2023.06.002.",
    "You, N., et al. SpliceTransformer predicts tissue-specific splicing linked to human diseases. Nature Communications, 2024. doi:10.1038/s41467-024-53088-6.",
    "Xiong, H.Y., et al. RNA splicing. The human splicing code reveals new insights into the genetic determinants of disease (SPANR). Science 347(6218):1254806, 2015. doi:10.1126/science.1254806.",
    "Cheng, J., Nguyen, T.Y.D., Cygan, K.J., et al. MMSplice: modular modeling improves the predictions of genetic variant effects on splicing. Genome Biology 20:48, 2019. doi:10.1186/s13059-019-1653-z.",
    "Rosenberg, A.B., Patwardhan, R.P., Shendure, J., and Seelig, G. Learning the sequence determinants of alternative splicing from millions of random sequences. Cell 163:698–711, 2015. doi:10.1016/j.cell.2015.09.054.",
    "Chong, R., Insigne, K.D., Yao, D., Burghard, C.P., Wang, J., Hsiao, Y.-H.E., Jones, E.M., Goodman, D.B., Xiao, X., and Kosuri, S. A multiplexed assay for exon recognition reveals that an unappreciated fraction of rare genetic variants cause large-effect splicing disruptions (MFASS). Molecular Cell 73:183–194.e8, 2019. doi:10.1016/j.molcel.2018.10.037.",
    "Riepe, T.V., et al. Benchmarking deep learning splice prediction tools using functional splice assays. Human Mutation 42:799–810, 2021. doi:10.1002/humu.24212.",
    "Merchant, A.T., et al. A high-level programming language for generative biology with Proto. bioRxiv 2026.06.22.733870, 2026. doi:10.64898/2026.06.22.733870.",
    "Saito, T. and Rehmsmeier, M. The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLOS ONE 10:e0118432, 2015. doi:10.1371/journal.pone.0118432."]:
    d.add_paragraph(ref, style="List Bullet")

d.save(str(OUT)); print("wrote", OUT)
